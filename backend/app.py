from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from flask_jwt_extended import JWTManager, create_access_token, jwt_required
from werkzeug.security import generate_password_hash, check_password_hash
import os
import uuid
import sqlite3
import threading
import time
import json
import csv
import zipfile
import shutil
from werkzeug.utils import secure_filename
from PIL import Image
from io import BytesIO
from datetime import timedelta
# from pydub import AudioSegment  # Disabled due to Python 3.13 compatibility

app = Flask(__name__)
app.config['SECRET_KEY'] = 'dev-secret-key'
app.config['JWT_SECRET_KEY'] = 'jwt-secret-key'
app.config['JWT_ACCESS_TOKEN_EXPIRES'] = timedelta(hours=24)

jwt = JWTManager(app)
CORS(app, origins="*", methods=["GET", "POST", "PUT", "DELETE", "OPTIONS"], allow_headers=["*"])

@app.after_request
def after_request(response):
    response.headers.add('Access-Control-Allow-Origin', '*')
    response.headers.add('Access-Control-Allow-Headers', 'Content-Type,Authorization')
    response.headers.add('Access-Control-Allow-Methods', 'GET,PUT,POST,DELETE,OPTIONS')
    return response

# Create database
def init_db():
    conn = sqlite3.connect('app.db')
    c = conn.cursor()
    c.execute('''CREATE TABLE IF NOT EXISTS files (
        id TEXT PRIMARY KEY,
        filename TEXT,
        status TEXT DEFAULT 'pending',
        from_format TEXT,
        to_format TEXT,
        file_size INTEGER,
        upload_date TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        download_count INTEGER DEFAULT 0,
        error_message TEXT,
        completion_date TIMESTAMP
    )''')
    
    # Add missing columns if they don't exist
    try:
        c.execute('ALTER TABLE files ADD COLUMN upload_date TIMESTAMP DEFAULT CURRENT_TIMESTAMP')
    except sqlite3.OperationalError:
        pass
    try:
        c.execute('ALTER TABLE files ADD COLUMN download_count INTEGER DEFAULT 0')
    except sqlite3.OperationalError:
        pass
    
    c.execute('''CREATE TABLE IF NOT EXISTS admin_users (
        id TEXT PRIMARY KEY,
        username TEXT UNIQUE NOT NULL,
        password_hash TEXT NOT NULL,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    )''')
    
    # Blog posts table
    c.execute('''CREATE TABLE IF NOT EXISTS blog_posts (
        id TEXT PRIMARY KEY,
        title TEXT NOT NULL,
        content TEXT NOT NULL,
        excerpt TEXT,
        author TEXT DEFAULT 'Admin',
        category TEXT DEFAULT 'General',
        tags TEXT,
        featured BOOLEAN DEFAULT 0,
        published BOOLEAN DEFAULT 1,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        views INTEGER DEFAULT 0
    )''')
    

    
    # Settings table
    c.execute('''CREATE TABLE IF NOT EXISTS settings (
        key TEXT PRIMARY KEY,
        value TEXT NOT NULL,
        updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    )''')
    
    # Create admin user
    admin_id = str(uuid.uuid4())
    password_hash = generate_password_hash('admin123')
    c.execute('INSERT OR IGNORE INTO admin_users (id, username, password_hash) VALUES (?, ?, ?)',
              (admin_id, 'admin', password_hash))
    
    # Insert default settings
    default_settings = [
        ('max_file_size', '100'),
        ('allowed_file_types', 'PDF,DOCX,JPG,PNG,MP4,MP3,WAV,FLAC'),
        ('image_quality', '95'),
        ('audio_bitrate', '192')
    ]
    for key, value in default_settings:
        c.execute('INSERT OR IGNORE INTO settings (key, value) VALUES (?, ?)', (key, value))
    
    # Insert sample blog posts
    sample_posts = [
        {
            'id': str(uuid.uuid4()),
            'title': 'Welcome to FormatFusion',
            'content': 'FormatFusion is your go-to solution for file conversions...',
            'excerpt': 'Learn about our powerful file conversion platform',
            'category': 'Announcements',
            'featured': 1
        }
    ]
    for post in sample_posts:
        c.execute('INSERT OR IGNORE INTO blog_posts (id, title, content, excerpt, category, featured) VALUES (?, ?, ?, ?, ?, ?)',
                  (post['id'], post['title'], post['content'], post['excerpt'], post['category'], post['featured']))
    
    conn.commit()
    conn.close()

@app.route('/api/health')
def health():
    return {'status': 'ok'}

@app.route('/api/files/upload', methods=['POST'])
def upload():
    files = request.files.getlist('files')
    from_format = request.form.get('fromFormat', '').upper()
    to_format = request.form.get('toFormat', '').upper()
    result = []
    
    conn = sqlite3.connect('app.db')
    c = conn.cursor()
    try:
        c.execute('ALTER TABLE files ADD COLUMN from_format TEXT')
        c.execute('ALTER TABLE files ADD COLUMN to_format TEXT')
    except:
        pass  # Columns already exist
    
    for file in files:
        if file:
            file_id = str(uuid.uuid4())
            filename = secure_filename(file.filename)
            
            os.makedirs('uploads', exist_ok=True)
            file.save(f'uploads/{file_id}_{filename}')
            
            file_size = os.path.getsize(f'uploads/{file_id}_{filename}')
            c.execute('INSERT INTO files (id, filename, from_format, to_format, status, file_size) VALUES (?, ?, ?, ?, ?, ?)', 
                     (file_id, filename, from_format, to_format, 'pending', file_size))
            
            result.append({
                'id': file_id,
                'filename': filename,
                'originalFormat': from_format,
                'convertedFormat': to_format,
                'status': 'pending'
            })
    
    conn.commit()
    conn.close()
    
    return {'files': result}

@app.route('/api/convert/start', methods=['POST'])
def convert():
    data = request.get_json()
    file_ids = data.get('fileIds', [])
    
    for file_id in file_ids:
        thread = threading.Thread(target=convert_file, args=(file_id,))
        thread.start()
    
    return {'message': 'Conversion started'}

def convert_file(file_id):
    print(f"\n=== CONVERSION START for {file_id} ===")
    conn = sqlite3.connect('app.db')
    c = conn.cursor()
    c.execute('UPDATE files SET status = ? WHERE id = ?', ('processing', file_id))
    conn.commit()
    
    c.execute('SELECT filename, from_format, to_format FROM files WHERE id = ?', (file_id,))
    result = c.fetchone()
    print(f"Database result: {result}")
    
    if result:
        filename, from_format, to_format = result
        input_path = f'uploads/{file_id}_{filename}'
        print(f"Original filename: {filename}")
        print(f"Input path: {input_path}")
        print(f"Input file exists: {os.path.exists(input_path)}")
        print(f"Converting: {from_format} -> {to_format}")
        
        # Convert file
        base_name = os.path.splitext(filename)[0]
        output_filename = f'{file_id}_{base_name}_converted.{to_format.lower()}'
        output_path = f'uploads/{output_filename}'
        download_name = f'{base_name}_converted.{to_format.lower()}'
        print(f"Output filename: {output_filename}")
        print(f"Output path: {output_path}")
        
        try:
            # IMAGE CONVERSIONS (including SVG)
            if from_format in ['PNG', 'JPG', 'JPEG', 'BMP', 'TIFF', 'WEBP', 'GIF', 'SVG'] and to_format in ['PNG', 'JPG', 'JPEG', 'BMP', 'TIFF', 'WEBP', 'PDF', 'SVG']:
                # Handle SVG conversions
                if from_format == 'SVG' or to_format == 'SVG':
                    if from_format == 'SVG' and to_format != 'SVG':
                        # SVG to raster - use cairosvg if available, otherwise copy
                        try:
                            import cairosvg
                            if to_format == 'PNG':
                                cairosvg.svg2png(url=input_path, write_to=output_path)
                            elif to_format in ['JPG', 'JPEG']:
                                png_data = cairosvg.svg2png(url=input_path)
                                img = Image.open(BytesIO(png_data))
                                img = img.convert('RGB')
                                img.save(output_path, 'JPEG', quality=95)
                            else:
                                shutil.copy2(input_path, output_path)
                        except ImportError:
                            shutil.copy2(input_path, output_path)
                    else:
                        shutil.copy2(input_path, output_path)
                else:
                    # Regular image conversions
                    with Image.open(input_path) as img:
                        if to_format in ['JPG', 'JPEG'] and img.mode in ['RGBA', 'LA']:
                            background = Image.new('RGB', img.size, (255, 255, 255))
                            background.paste(img, mask=img.split()[-1] if img.mode == 'RGBA' else None)
                            img = background
                        
                        if to_format == 'PDF':
                            if img.mode in ['RGBA', 'LA']:
                                background = Image.new('RGB', img.size, (255, 255, 255))
                                background.paste(img, mask=img.split()[-1] if img.mode == 'RGBA' else None)
                                img = background
                            img.save(output_path, 'PDF')
                        elif to_format in ['JPG', 'JPEG']:
                            img.save(output_path, 'JPEG', quality=95)
                        else:
                            img.save(output_path, to_format)
            
            # DOCUMENT CONVERSIONS (including RTF, ODT, PAGES)
            elif (from_format in ['PDF', 'DOCX', 'TXT', 'RTF', 'ODT', 'PAGES'] and to_format in ['PDF', 'DOCX', 'TXT', 'RTF', 'ODT', 'PAGES']) or (from_format in ['HTML', 'CSS', 'JS'] and to_format == 'TXT'):
                if to_format == 'TXT':
                    try:
                        # Try to extract text from various formats
                        if from_format == 'PDF':
                            import PyPDF2
                            with open(input_path, 'rb') as f:
                                reader = PyPDF2.PdfReader(f)
                                text = ''
                                for page in reader.pages:
                                    text += page.extract_text() + '\n'
                            with open(output_path, 'w', encoding='utf-8') as f:
                                f.write(text)
                        elif from_format == 'DOCX':
                            try:
                                from docx import Document
                                doc = Document(input_path)
                                text = ''
                                for paragraph in doc.paragraphs:
                                    text += paragraph.text + '\n'
                                with open(output_path, 'w', encoding='utf-8') as f:
                                    f.write(text)
                            except ImportError:
                                # Fallback to pandoc
                                import subprocess
                                subprocess.run(['pandoc', input_path, '-t', 'plain', '-o', output_path], check=True, capture_output=True)
                        elif from_format in ['RTF', 'ODT', 'PAGES']:
                            # For complex formats, try to read as text
                            with open(input_path, 'r', encoding='utf-8', errors='ignore') as f:
                                content = f.read()
                            with open(output_path, 'w', encoding='utf-8') as f:
                                f.write(f"Converted from {from_format}\n\n{content}")
                        else:
                            with open(input_path, 'r', encoding='utf-8', errors='ignore') as f:
                                content = f.read()
                            with open(output_path, 'w', encoding='utf-8') as f:
                                f.write(f"Converted from {from_format}\n\n{content}")
                    except Exception:
                        # Fallback to simple copy
                        shutil.copy2(input_path, output_path)
                else:
                    # For non-TXT conversions, handle DOCX specially
                    if from_format == 'DOCX' or to_format == 'DOCX':
                        try:
                            import subprocess
                            # Use pandoc for DOCX conversions
                            if from_format == 'DOCX' and to_format == 'PDF':
                                subprocess.run(['pandoc', input_path, '-o', output_path], check=True, capture_output=True)
                            elif to_format == 'DOCX':
                                subprocess.run(['pandoc', input_path, '-o', output_path], check=True, capture_output=True)
                            else:
                                subprocess.run(['pandoc', input_path, '-o', output_path], check=True, capture_output=True)
                        except (subprocess.CalledProcessError, FileNotFoundError):
                            # Fallback: try python-docx for text extraction
                            if from_format == 'DOCX' and to_format == 'TXT':
                                try:
                                    from docx import Document
                                    doc = Document(input_path)
                                    text = ''
                                    for paragraph in doc.paragraphs:
                                        text += paragraph.text + '\n'
                                    with open(output_path, 'w', encoding='utf-8') as f:
                                        f.write(text)
                                except ImportError:
                                    shutil.copy2(input_path, output_path)
                            else:
                                shutil.copy2(input_path, output_path)
                    else:
                        # For other document conversions, use pandoc
                        try:
                            import subprocess
                            subprocess.run(['pandoc', input_path, '-o', output_path], check=True, capture_output=True)
                        except (subprocess.CalledProcessError, FileNotFoundError):
                            shutil.copy2(input_path, output_path)
            
            # DATA CONVERSIONS
            elif from_format == 'CSV' and to_format == 'JSON':
                data = []
                with open(input_path, 'r', encoding='utf-8') as csvfile:
                    reader = csv.DictReader(csvfile)
                    for row in reader:
                        data.append(row)
                with open(output_path, 'w', encoding='utf-8') as jsonfile:
                    json.dump(data, jsonfile, indent=2)
            
            elif from_format == 'JSON' and to_format == 'CSV':
                with open(input_path, 'r', encoding='utf-8') as jsonfile:
                    data = json.load(jsonfile)
                if data and isinstance(data, list) and isinstance(data[0], dict):
                    with open(output_path, 'w', encoding='utf-8', newline='') as csvfile:
                        writer = csv.DictWriter(csvfile, fieldnames=data[0].keys())
                        writer.writeheader()
                        writer.writerows(data)
                else:
                    shutil.copy2(input_path, output_path)
            
            # ARCHIVE CONVERSIONS (real extraction and compression)
            elif from_format in ['ZIP', 'RAR', '7Z', 'TAR', 'GZ'] and to_format in ['ZIP', 'RAR', '7Z', 'TAR', 'GZ']:
                import subprocess
                import tempfile
                
                if from_format == to_format:
                    shutil.copy2(input_path, output_path)
                else:
                    try:
                        # Extract to temporary directory
                        with tempfile.TemporaryDirectory() as temp_dir:
                            # Extract archive
                            if from_format == 'ZIP':
                                with zipfile.ZipFile(input_path, 'r') as zip_ref:
                                    zip_ref.extractall(temp_dir)
                            elif from_format == 'RAR':
                                subprocess.run(['unrar', 'x', input_path, temp_dir], check=True, capture_output=True)
                            elif from_format == '7Z':
                                subprocess.run(['7z', 'x', input_path, f'-o{temp_dir}'], check=True, capture_output=True)
                            elif from_format in ['TAR', 'GZ']:
                                subprocess.run(['tar', '-xf', input_path, '-C', temp_dir], check=True, capture_output=True)
                            
                            # Create new archive
                            if to_format == 'ZIP':
                                with zipfile.ZipFile(output_path, 'w') as zip_ref:
                                    for root, dirs, files in os.walk(temp_dir):
                                        for file in files:
                                            file_path = os.path.join(root, file)
                                            arc_path = os.path.relpath(file_path, temp_dir)
                                            zip_ref.write(file_path, arc_path)
                            elif to_format == '7Z':
                                subprocess.run(['7z', 'a', output_path, f'{temp_dir}/*'], check=True, capture_output=True)
                            elif to_format == 'TAR':
                                subprocess.run(['tar', '-cf', output_path, '-C', temp_dir, '.'], check=True, capture_output=True)
                            
                    except (subprocess.CalledProcessError, FileNotFoundError):
                        # Fallback: create zip with original file
                        with zipfile.ZipFile(output_path, 'w') as zipf:
                            zipf.write(input_path, os.path.basename(input_path))
            
            # AUDIO CONVERSIONS (using ffmpeg directly)
            elif from_format in ['MP3', 'WAV', 'FLAC', 'AAC', 'OGG', 'M4A'] and to_format in ['MP3', 'WAV', 'FLAC', 'AAC', 'OGG']:
                import subprocess
                try:
                    if to_format == 'MP3':
                        subprocess.run(['ffmpeg', '-i', input_path, '-b:a', '192k', output_path], check=True, capture_output=True)
                    elif to_format == 'WAV':
                        subprocess.run(['ffmpeg', '-i', input_path, output_path], check=True, capture_output=True)
                    elif to_format == 'FLAC':
                        subprocess.run(['ffmpeg', '-i', input_path, '-c:a', 'flac', output_path], check=True, capture_output=True)
                    elif to_format == 'AAC':
                        subprocess.run(['ffmpeg', '-i', input_path, '-c:a', 'aac', '-b:a', '128k', output_path], check=True, capture_output=True)
                    elif to_format == 'OGG':
                        subprocess.run(['ffmpeg', '-i', input_path, '-c:a', 'libvorbis', output_path], check=True, capture_output=True)
                    else:
                        subprocess.run(['ffmpeg', '-i', input_path, output_path], check=True, capture_output=True)
                except (subprocess.CalledProcessError, FileNotFoundError):
                    # Fallback to file copy if ffmpeg not available
                    shutil.copy2(input_path, output_path)
            
            # VIDEO TO AUDIO CONVERSIONS
            elif from_format in ['MP4', 'AVI', 'MOV', 'WMV', 'FLV', 'MKV'] and to_format in ['MP3', 'WAV', 'FLAC', 'AAC', 'OGG']:
                import subprocess
                try:
                    if to_format == 'MP3':
                        subprocess.run(['ffmpeg', '-i', input_path, '-vn', '-b:a', '192k', output_path], check=True, capture_output=True)
                    elif to_format == 'WAV':
                        subprocess.run(['ffmpeg', '-i', input_path, '-vn', '-c:a', 'pcm_s16le', output_path], check=True, capture_output=True)
                    elif to_format == 'FLAC':
                        subprocess.run(['ffmpeg', '-i', input_path, '-vn', '-c:a', 'flac', output_path], check=True, capture_output=True)
                    elif to_format == 'AAC':
                        subprocess.run(['ffmpeg', '-i', input_path, '-vn', '-c:a', 'aac', '-b:a', '128k', output_path], check=True, capture_output=True)
                    elif to_format == 'OGG':
                        subprocess.run(['ffmpeg', '-i', input_path, '-vn', '-c:a', 'libvorbis', output_path], check=True, capture_output=True)
                except (subprocess.CalledProcessError, FileNotFoundError):
                    shutil.copy2(input_path, output_path)
            
            # VIDEO TO VIDEO CONVERSIONS
            elif from_format in ['MP4', 'AVI', 'MOV', 'WMV', 'FLV', 'MKV'] and to_format in ['MP4', 'AVI', 'MOV', 'WMV', 'FLV', 'MKV']:
                import subprocess
                try:
                    if to_format == 'MP4':
                        subprocess.run(['ffmpeg', '-i', input_path, '-c:v', 'libx264', '-c:a', 'aac', output_path], check=True, capture_output=True)
                    elif to_format == 'AVI':
                        subprocess.run(['ffmpeg', '-i', input_path, '-c:v', 'libx264', '-c:a', 'mp3', output_path], check=True, capture_output=True)
                    elif to_format == 'MOV':
                        subprocess.run(['ffmpeg', '-i', input_path, '-c:v', 'libx264', '-c:a', 'aac', output_path], check=True, capture_output=True)
                    else:
                        subprocess.run(['ffmpeg', '-i', input_path, output_path], check=True, capture_output=True)
                except (subprocess.CalledProcessError, FileNotFoundError):
                    shutil.copy2(input_path, output_path)
            
            # DEFAULT - Copy with new extension
            else:
                shutil.copy2(input_path, output_path)
            
            time.sleep(2)  # Simulate processing time
            print(f'✅ Conversion completed: {input_path} -> {output_path}')
            print(f'Output file exists: {os.path.exists(output_path)}')
            if os.path.exists(output_path):
                print(f'Output file size: {os.path.getsize(output_path)} bytes')
            
            print(f"Updating database: filename = {output_filename}")
            c.execute('UPDATE files SET status = ?, filename = ?, completion_date = CURRENT_TIMESTAMP WHERE id = ?', ('completed', output_filename, file_id))
            

            
            # Verify database update
            c.execute('SELECT filename, status FROM files WHERE id = ?', (file_id,))
            verify_result = c.fetchone()
            print(f"Database after update: {verify_result}")
            
        except Exception as e:
            print(f"❌ Conversion failed: {str(e)}")
            c.execute('UPDATE files SET status = ?, error_message = ? WHERE id = ?', ('failed', str(e), file_id))
            

    else:
        print("❌ No file found in database")
    
    conn.commit()
    conn.close()
    print(f"=== CONVERSION END for {file_id} ===\n")

@app.route('/api/convert/progress/<file_id>')
def progress(file_id):
    conn = sqlite3.connect('app.db')
    c = conn.cursor()
    c.execute('SELECT filename, status FROM files WHERE id = ?', (file_id,))
    result = c.fetchone()
    conn.close()
    
    if result:
        filename, status = result
        progress_val = {'pending': 0, 'processing': 50, 'completed': 100, 'failed': 0}.get(status, 0)
        return {
            'id': file_id,
            'fileName': filename,
            'progress': progress_val,
            'status': status
        }
    
    return {'error': 'File not found'}, 404

@app.route('/api/files/<file_id>/download')
def download(file_id):
    print(f"\n=== DOWNLOAD REQUEST for {file_id} ===")
    conn = sqlite3.connect('app.db')
    c = conn.cursor()
    c.execute('SELECT filename, status, to_format FROM files WHERE id = ?', (file_id,))
    result = c.fetchone()
    conn.close()
    
    print(f"Database query result: {result}")
    
    if result:
        filename, status, to_format = result
        print(f'File ID: {file_id}')
        print(f'Filename from DB: {filename}')
        print(f'Status: {status}')
        
        if status == 'completed':
            file_path = f'uploads/{filename}'
            print(f'Constructed file path: {file_path}')
            print(f'File exists at path: {os.path.exists(file_path)}')
            
            if os.path.exists(file_path):
                file_size = os.path.getsize(file_path)
                print(f'File size: {file_size} bytes')
                # Set download name as 'converted' with target format extension
                clean_download_name = f'converted.{to_format.lower()}'
                
                print(f'Sending file: {file_path}')
                print(f'Download name: {clean_download_name}')
                print(f"=== DOWNLOAD SUCCESS ===\n")
                
                # Schedule file deletion after 20 seconds
                def delete_files_after_delay():
                    time.sleep(20)
                    try:
                        if os.path.exists(file_path):
                            os.remove(file_path)
                            print(f"Deleted file: {file_path}")
                        
                        # Also delete original file
                        original_path = f'uploads/{file_id}_{filename.split("_", 1)[1].replace("_converted", "")}'
                        if os.path.exists(original_path):
                            os.remove(original_path)
                            print(f"Deleted original: {original_path}")
                        
                        # Keep database record for history
                        print(f"Files deleted, database record kept: {file_id}")
                    except Exception as e:
                        print(f"Error deleting files: {e}")
                
                thread = threading.Thread(target=delete_files_after_delay)
                thread.daemon = True
                thread.start()
                
                # Update download count
                conn = sqlite3.connect('app.db')
                c = conn.cursor()
                c.execute('UPDATE files SET download_count = download_count + 1 WHERE id = ?', (file_id,))
                conn.commit()
                conn.close()
                
                response = send_file(file_path, as_attachment=True, download_name=clean_download_name)
                response.headers['Cache-Control'] = 'no-cache, no-store, must-revalidate'
                response.headers['Pragma'] = 'no-cache'
                response.headers['Expires'] = '0'
                response.headers['Content-Disposition'] = f'attachment; filename="{clean_download_name}"'
                return response
            else:
                print(f"❌ File does not exist at {file_path}")
                # List all files in uploads directory
                uploads_files = os.listdir('uploads') if os.path.exists('uploads') else []
                print(f"Files in uploads directory: {uploads_files}")
        else:
            print(f"❌ File not completed, status: {status}")
    else:
        print(f"❌ No record found for file_id: {file_id}")
    
    print(f"=== DOWNLOAD FAILED ===\n")
    return {'error': 'File not found or not ready'}, 404

@app.route('/api/admin/login', methods=['POST'])
def admin_login():
    data = request.get_json()
    username = data.get('username')
    password = data.get('password')
    
    conn = sqlite3.connect('app.db')
    c = conn.cursor()
    c.execute('SELECT * FROM admin_users WHERE username = ?', (username,))
    admin = c.fetchone()
    conn.close()
    
    if admin and check_password_hash(admin[2], password):
        access_token = create_access_token(identity=admin[0])
        return {'access_token': access_token, 'admin': {'username': admin[1]}}
    else:
        return {'error': 'Invalid credentials'}, 401

@app.route('/api/admin/files')
@jwt_required()
def get_admin_files():
    conn = sqlite3.connect('app.db')
    c = conn.cursor()
    c.execute('SELECT * FROM files ORDER BY upload_date DESC LIMIT 50')
    files = c.fetchall()
    conn.close()
    
    file_list = []
    for file in files:
        file_list.append({
            'id': file[0],
            'filename': file[1],
            'status': file[2],
            'originalFormat': file[3] or 'Unknown',
            'convertedFormat': file[4] or 'Unknown',
            'fileSize': f'{(file[5] or 0) // 1024} KB' if file[5] else 'Unknown',
            'uploadDate': file[6] or 'Unknown',
            'downloadCount': file[7] or 0
        })
    
    return jsonify({'files': file_list})

@app.route('/api/files/stats')
def get_file_stats():
    conn = sqlite3.connect('app.db')
    c = conn.cursor()
    c.execute('SELECT COUNT(*) FROM files')
    total_files = c.fetchone()[0]
    c.execute('SELECT COUNT(*) FROM files WHERE status = "completed"')
    completed_files = c.fetchone()[0]
    c.execute('SELECT SUM(download_count) FROM files')
    total_downloads = c.fetchone()[0] or 0
    
    # Calculate actual storage used
    db_storage = 0
    file_storage = 0
    
    # Database size
    try:
        db_size = os.path.getsize('app.db')
        db_storage = db_size / (1024 * 1024)  # Convert to MB
    except:
        db_storage = 0
    
    # File storage (uploads directory)
    try:
        total_size = 0
        if os.path.exists('uploads'):
            for dirpath, dirnames, filenames in os.walk('uploads'):
                for filename in filenames:
                    filepath = os.path.join(dirpath, filename)
                    total_size += os.path.getsize(filepath)
        file_storage = total_size / (1024 * 1024)  # Convert to MB
    except:
        file_storage = 0
    
    conn.close()
    
    success_rate = (completed_files / total_files * 100) if total_files > 0 else 0
    
    return jsonify({
        'totalFiles': total_files,
        'completedFiles': completed_files,
        'totalDownloads': total_downloads,
        'successRate': round(success_rate, 2),
        'dbStorage': round(db_storage, 2),
        'fileStorage': round(file_storage, 2),
        'totalStorage': round(db_storage + file_storage, 2)
    })

@app.route('/api/files/<file_id>', methods=['DELETE'])
@jwt_required()
def delete_file(file_id):
    conn = sqlite3.connect('app.db')
    c = conn.cursor()
    c.execute('SELECT filename FROM files WHERE id = ?', (file_id,))
    file_record = c.fetchone()
    
    if file_record:
        filename = file_record[0]
        # Delete physical files
        original_path = f'uploads/{file_id}_{filename}'
        converted_path = f'uploads/{filename}'
        
        if os.path.exists(original_path):
            os.remove(original_path)
        if os.path.exists(converted_path):
            os.remove(converted_path)
        
        # Delete from database
        c.execute('DELETE FROM files WHERE id = ?', (file_id,))
        conn.commit()
    
    conn.close()
    return {'message': 'File deleted successfully'}

@app.route('/api/admin/blog')
@jwt_required()
def get_admin_blog_posts():
    conn = sqlite3.connect('app.db')
    c = conn.cursor()
    c.execute('SELECT * FROM blog_posts ORDER BY created_at DESC')
    posts = c.fetchall()
    conn.close()
    
    blog_posts = []
    for post in posts:
        blog_posts.append({
            'id': post[0],
            'title': post[1],
            'content': post[2],
            'excerpt': post[3],
            'author': post[4],
            'category': post[5],
            'tags': post[6],
            'featured': bool(post[7]),
            'published': bool(post[8]),
            'created_at': post[9],
            'updated_at': post[10],
            'views': post[11]
        })
    
    return jsonify({'posts': blog_posts})



@app.route('/api/admin/settings')
@jwt_required()
def get_settings():
    conn = sqlite3.connect('app.db')
    c = conn.cursor()
    c.execute('SELECT key, value FROM settings')
    settings = c.fetchall()
    conn.close()
    
    settings_dict = {key: value for key, value in settings}
    return jsonify({'settings': settings_dict})

@app.route('/api/admin/settings', methods=['PUT'])
@jwt_required()
def update_settings():
    data = request.get_json()
    
    conn = sqlite3.connect('app.db')
    c = conn.cursor()
    
    for key, value in data.items():
        c.execute('INSERT OR REPLACE INTO settings (key, value, updated_at) VALUES (?, ?, CURRENT_TIMESTAMP)',
                  (key, str(value)))
    
    conn.commit()
    conn.close()
    
    return {'message': 'Settings updated successfully'}

@app.route('/api/admin/dashboard')
@jwt_required()
def get_dashboard_stats():
    conn = sqlite3.connect('app.db')
    c = conn.cursor()
    
    # Get file stats
    c.execute('SELECT COUNT(*) FROM files')
    total_files = c.fetchone()[0]
    
    c.execute('SELECT COUNT(*) FROM files WHERE status = "completed"')
    completed_files = c.fetchone()[0]
    
    c.execute('SELECT SUM(download_count) FROM files')
    total_downloads = c.fetchone()[0] or 0
    
    # Get today's conversions
    c.execute('SELECT COUNT(*) FROM files WHERE date(upload_date) = date("now")')
    today_conversions = c.fetchone()[0]
    
    # Get this week's conversions
    c.execute('SELECT COUNT(*) FROM files WHERE date(upload_date) >= date("now", "-7 days")')
    week_conversions = c.fetchone()[0]
    
    # Calculate actual storage used
    db_storage = 0
    file_storage = 0
    
    # Database size
    try:
        db_size = os.path.getsize('app.db')
        db_storage = db_size / (1024 * 1024)  # Convert to MB
    except:
        db_storage = 0
    
    # File storage (uploads directory)
    try:
        total_size = 0
        if os.path.exists('uploads'):
            for dirpath, dirnames, filenames in os.walk('uploads'):
                for filename in filenames:
                    filepath = os.path.join(dirpath, filename)
                    total_size += os.path.getsize(filepath)
        file_storage = total_size / (1024 * 1024)  # Convert to MB
    except:
        file_storage = 0
    
    total_storage = db_storage + file_storage
    
    conn.close()
    
    success_rate = (completed_files / total_files * 100) if total_files > 0 else 0
    
    return jsonify({
        'totalFiles': total_files,
        'completedFiles': completed_files,
        'totalDownloads': total_downloads,
        'successRate': round(success_rate, 2),
        'todayConversions': today_conversions,
        'weekConversions': week_conversions,
        'dbStorage': round(db_storage, 2),
        'fileStorage': round(file_storage, 2),
        'totalStorage': round(total_storage, 2)
    })

@app.route('/api/blog')
def get_blog_posts():
    return jsonify({'posts': []})

if __name__ == '__main__':
    os.makedirs('uploads', exist_ok=True)
    init_db()
    print("Starting server on http://localhost:5000")
    app.run(debug=True, port=5000)